# backend/app/services/document.py
import io
import json
from typing import Tuple, Dict, Any
import docx
from PyPDF2 import PdfReader
from bs4 import BeautifulSoup

async def process_document(filename: str, file_content: bytes) -> Tuple[str, str]:
    """
    Process an uploaded document and extract its text content
    Returns a tuple of (file_type, extracted_text)
    """
    # Determine file type by extension
    extension = filename.lower().split('.')[-1]
    
    if extension == 'pdf':
        return process_pdf(file_content)
    elif extension in ['docx', 'doc']:
        return process_docx(file_content)
    elif extension == 'html':
        return process_html(file_content)
    elif extension == 'json':
        return process_json(file_content)
    elif extension in ['txt', 'text']:
        return process_text(file_content)
    else:
        # Unsupported file type
        raise ValueError(f"Unsupported file type: {extension}")

def process_pdf(file_content: bytes) -> Tuple[str, str]:
    """Extract text from PDF file"""
    pdf_file = io.BytesIO(file_content)
    reader = PdfReader(pdf_file)
    
    text_content = []
    for page in reader.pages:
        text_content.append(page.extract_text())
    
    return "pdf", "\n\n".join(text_content)

def process_docx(file_content: bytes) -> Tuple[str, str]:
    """Extract text from DOCX file"""
    docx_file = io.BytesIO(file_content)
    doc = docx.Document(docx_file)
    
    text_content = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            text_content.append(paragraph.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text:
                    row_text.append(cell.text)
            if row_text:
                text_content.append(" | ".join(row_text))
    
    return "docx", "\n\n".join(text_content)

def process_html(file_content: bytes) -> Tuple[str, str]:
    """Extract text from HTML file"""
    try:
        # Try to decode as UTF-8
        html_content = file_content.decode('utf-8')
    except UnicodeDecodeError:
        # Fall back to latin-1 if UTF-8 fails
        html_content = file_content.decode('latin-1')
    
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Remove script and style elements
    for script in soup(["script", "style"]):
        script.extract()
    
    # Get text
    text = soup.get_text(separator='\n')
    
    # Break into lines and remove leading/trailing space
    lines = (line.strip() for line in text.splitlines())
    # Break multi-headlines into a line each
    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
    # Drop blank lines
    text_content = '\n'.join(chunk for chunk in chunks if chunk)
    
    return "html", text_content

def process_json(file_content: bytes) -> Tuple[str, str]:
    """Extract text from JSON file"""
    try:
        # Try to decode as UTF-8
        json_str = file_content.decode('utf-8')
        # Parse JSON to validate format
        json_data = json.loads(json_str)
        
        # Convert to string representation
        if isinstance(json_data, Dict):
            # If it's a dict, try to extract content intelligently
            text_content = extract_json_content(json_data)
        else:
            # Otherwise just use the string representation
            text_content = json_str
            
        return "json", text_content
    except Exception as e:
        raise ValueError(f"Invalid JSON file: {str(e)}")

def extract_json_content(json_data: Dict[str, Any]) -> str:
    """
    Attempt to intelligently extract text content from JSON
    Looks for known fields like 'content', 'text', 'description', etc.
    """
    # Look for content fields
    content_fields = ['content', 'text', 'body', 'description', 'data']
    
    for field in content_fields:
        if field in json_data and isinstance(json_data[field], str):
            return json_data[field]
    
    # If we have items/elements, extract content from each
    if 'items' in json_data and isinstance(json_data['items'], list):
        return '\n\n'.join(extract_json_content(item) if isinstance(item, dict) else str(item) 
                           for item in json_data['items'])
    
    if 'elements' in json_data and isinstance(json_data['elements'], list):
        return '\n\n'.join(extract_json_content(item) if isinstance(item, dict) else str(item) 
                           for item in json_data['elements'])
    
    # Fallback: serialize the whole JSON structure
    return json.dumps(json_data)

def process_text(file_content: bytes) -> Tuple[str, str]:
    """Extract text from plain text file"""
    try:
        # Try to decode as UTF-8
        text = file_content.decode('utf-8')
    except UnicodeDecodeError:
        # Fall back to latin-1 if UTF-8 fails
        text = file_content.decode('latin-1')
    
    return "text", text